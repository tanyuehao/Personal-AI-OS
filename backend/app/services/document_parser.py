"""
Personal AI OS - Document Parser
文档解析服务
"""
import os
from typing import Optional
from abc import ABC, abstractmethod


class BaseParser(ABC):
    """文档解析器基类"""
    
    @abstractmethod
    async def parse(self, file_path: str) -> str:
        """解析文档并返回文本内容"""
        pass


class PDFParser(BaseParser):
    """PDF 文档解析器"""
    
    async def parse(self, file_path: str) -> str:
        """解析 PDF 文档"""
        try:
            import PyPDF2
            
            text_content = []
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text_content.append(page.extract_text())
            
            return "\n\n".join(text_content)
        except Exception as e:
            raise Exception(f"PDF 解析失败: {str(e)}")


class DocxParser(BaseParser):
    """Word 文档解析器"""
    
    async def parse(self, file_path: str) -> str:
        """解析 Word 文档"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            text_content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # 处理表格
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    text_content.append(" | ".join(row_text))
            
            return "\n\n".join(text_content)
        except Exception as e:
            raise Exception(f"Word 文档解析失败: {str(e)}")


class MarkdownParser(BaseParser):
    """Markdown 文档解析器"""
    
    async def parse(self, file_path: str) -> str:
        """解析 Markdown 文档"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            raise Exception(f"Markdown 解析失败: {str(e)}")


class TextParser(BaseParser):
    """纯文本文档解析器"""
    
    async def parse(self, file_path: str) -> str:
        """解析纯文本文档"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            raise Exception(f"文本解析失败: {str(e)}")


class CSVParser(BaseParser):
    """CSV 文档解析器"""
    
    async def parse(self, file_path: str) -> str:
        """解析 CSV 文档"""
        try:
            import csv
            
            rows = []
            with open(file_path, 'r', encoding='utf-8') as file:
                csv_reader = csv.reader(file)
                for row in csv_reader:
                    rows.append(" | ".join(row))
            
            return "\n".join(rows)
        except Exception as e:
            raise Exception(f"CSV 解析失败: {str(e)}")


class ExcelParser(BaseParser):
    """Excel 文档解析器"""
    
    async def parse(self, file_path: str) -> str:
        """解析 Excel 文档"""
        try:
            import openpyxl
            
            wb = openpyxl.load_workbook(file_path)
            text_content = []
            
            for sheet in wb.sheetnames:
                ws = wb[sheet]
                text_content.append(f"=== Sheet: {sheet} ===")
                
                for row in ws.iter_rows(values_only=True):
                    row_text = [str(cell) if cell is not None else "" for cell in row]
                    text_content.append(" | ".join(row_text))
            
            return "\n".join(text_content)
        except Exception as e:
            raise Exception(f"Excel 解析失败: {str(e)}")


def get_parser(file_type: str) -> BaseParser:
    """
    根据文件类型获取对应的解析器
    
    Args:
        file_type: 文件扩展名（如 .pdf, .docx 等）
    
    Returns:
        对应的文档解析器实例
    """
    parsers = {
        '.pdf': PDFParser(),
        '.doc': DocxParser(),
        '.docx': DocxParser(),
        '.md': MarkdownParser(),
        '.markdown': MarkdownParser(),
        '.txt': TextParser(),
        '.text': TextParser(),
        '.csv': CSVParser(),
        '.xlsx': ExcelParser(),
        '.xls': ExcelParser(),
    }
    
    parser = parsers.get(file_type.lower())
    if not parser:
        raise ValueError(f"不支持的文件类型: {file_type}")
    
    return parser


async def parse_document(file_path: str, file_type: str) -> str:
    """
    解析文档
    
    Args:
        file_path: 文件路径
        file_type: 文件类型
    
    Returns:
        解析后的文本内容
    """
    parser = get_parser(file_type)
    return await parser.parse(file_path)
